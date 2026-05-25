import os
import re
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set inner padding for a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_table_borders(table):
    """Apply elegant academic borders (top/bottom thick borders, thin header bottom border)."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="12" w:space="0" w:color="0F172A"/>'
        f'<w:bottom w:val="single" w:sz="12" w:space="0" w:color="0F172A"/>'
        f'<w:left w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        f'<w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def format_run(run, bold=False, italic=False, font_size=12, font_name="Times New Roman", color_rgb=(0,0,0)):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(*color_rgb)

def add_styled_paragraph(doc, text, style='Normal', space_after=6, line_spacing=1.15, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    p.alignment = align
    
    # Simple markdown parser for inline bold/italic
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            format_run(run, bold=True)
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            format_run(run, italic=True)
        else:
            run = p.add_run(part)
            format_run(run)
    return p

def main():
    report_md_path = r"f:\OneDrive\Phuong_2025\VIN\NLP\Project\report\report.md"
    report_docx_path = r"f:\OneDrive\Phuong_2025\VIN\NLP\Project\report\report.docx"
    
    print(f"Reading report markdown from: {report_md_path}")
    with open(report_md_path, "r", encoding="utf-8") as f:
        content = f.read()

    doc = Document()
    
    # Configure Normal Margins (~2.5cm / 1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    lines = content.split('\n')
    
    # Simple state machine for parsing Markdown structure
    in_code_block = False
    table_lines = []
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Code block handling
        if line.startswith("```"):
            in_code_block = not in_code_block
            i += 1
            continue
        
        if in_code_block:
            p = doc.add_paragraph(style='Normal')
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(lines[i])
            format_run(run, font_name="Consolas", font_size=10, color_rgb=(100, 110, 120))
            i += 1
            continue

        # Horizontal rule
        if line == "---":
            # Add a visual page break if it's the title page split, else a spacing
            if i < 30: # Near the top, after cover page info
                doc.add_page_break()
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(12)
                run = p.add_run("__________________________________________________________________")
                format_run(run, color_rgb=(200, 200, 200), font_size=10)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Headings
        if line.startswith("# "):
            title_text = line[2:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(title_text)
            format_run(run, bold=True, font_size=16, color_rgb=(15, 23, 42)) # Primary Deep Slate
            i += 1
            continue
            
        if line.startswith("## "):
            heading_text = line[3:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(heading_text)
            format_run(run, bold=True, font_size=14, color_rgb=(14, 116, 144)) # Accent Cyan/Teal
            i += 1
            continue
            
        if line.startswith("### "):
            subheading_text = line[4:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(subheading_text)
            format_run(run, bold=True, font_size=12, color_rgb=(124, 58, 237)) # Secondary Purple
            i += 1
            continue

        # Tables (lines starting with '|')
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            
            if len(table_lines) >= 2:
                # Parse header
                headers = [c.strip() for c in table_lines[0].split('|')[1:-1]]
                
                # Check for divider line
                start_row = 1
                if table_lines[1].replace('-','').replace('|','').replace(':','').strip() == "":
                    start_row = 2
                
                rows_data = []
                for row_line in table_lines[start_row:]:
                    row_cells = [c.strip() for c in row_line.split('|')[1:-1]]
                    rows_data.append(row_cells)
                
                # Add table to doc
                num_cols = len(headers)
                table = doc.add_table(rows=1, cols=num_cols)
                table.autofit = True
                set_table_borders(table)
                
                # Format headers
                hdr_cells = table.rows[0].cells
                for col_idx, text in enumerate(headers):
                    hdr_cells[col_idx].text = ""
                    p = hdr_cells[col_idx].paragraphs[0]
                    p.paragraph_format.space_after = Pt(2)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(text)
                    format_run(run, bold=True, font_size=10, color_rgb=(255, 255, 255))
                    set_cell_margins(hdr_cells[col_idx], top=120, bottom=120, left=150, right=150)
                    set_cell_shading(hdr_cells[col_idx], "0F172A") # Primary Deep Slate header background
                
                # Format rows
                for r_idx, row_data in enumerate(rows_data):
                    row_cells = table.add_row().cells
                    for col_idx in range(min(num_cols, len(row_data))):
                        text = row_data[col_idx]
                        row_cells[col_idx].text = ""
                        p = row_cells[col_idx].paragraphs[0]
                        p.paragraph_format.space_after = Pt(2)
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Apply cell margins
                        set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=120, right=120)
                        
                        # Strip HTML tags like <br> and replace with newline
                        clean_text = text.replace("<br>", "\n").replace("<br/>", "\n")
                        
                        run = p.add_run(clean_text)
                        format_run(run, font_size=10)
                        
                        # Alternating background colors
                        if r_idx % 2 == 1:
                            set_cell_shading(row_cells[col_idx], "F8FAFC")
            continue

        # Bullet lists
        if line.startswith("* ") or line.startswith("- "):
            list_text = line[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', list_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    format_run(run, bold=True)
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    format_run(run, italic=True)
                else:
                    run = p.add_run(part)
                    format_run(run)
            i += 1
            continue

        # Numbered lists (1. 2. etc)
        match_numbered = re.match(r'^(\d+)\.\s(.*)$', line)
        if match_numbered:
            num = match_numbered.group(1)
            list_text = match_numbered.group(2)
            p = doc.add_paragraph(style='Normal')
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.line_spacing = 1.15
            
            run_num = p.add_run(f"{num}. ")
            format_run(run_num, bold=True)
            
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', list_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    format_run(run, bold=True)
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    format_run(run, italic=True)
                else:
                    run = p.add_run(part)
                    format_run(run)
            i += 1
            continue

        # Image handling
        match_img = re.match(r'^!\[(.*?)\]\((.*?)\)$', line)
        if match_img:
            caption_text = match_img.group(1)
            image_name = match_img.group(2)
            img_path = os.path.join(r"f:\OneDrive\Phuong_2025\VIN\NLP\Project", image_name)
            
            if os.path.exists(img_path):
                print(f"Embedding picture: {img_path}")
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(8)
                p_img.paragraph_format.space_after = Pt(2)
                run_img = p_img.add_run()
                
                width_inches = 5.0 if "figure1" in image_name else 4.0
                run_img.add_picture(img_path, width=Inches(width_inches))
                
                # Add caption
                p_cap = doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_cap.paragraph_format.space_before = Pt(2)
                p_cap.paragraph_format.space_after = Pt(12)
                p_cap.paragraph_format.keep_with_next = True
                run_cap = p_cap.add_run(caption_text)
                format_run(run_cap, italic=True, font_size=10, color_rgb=(100, 110, 120))
            else:
                print(f"Warning: Image {image_name} not found.")
            i += 1
            continue

        # Regular paragraphs
        if line:
            align = WD_ALIGN_PARAGRAPH.LEFT
            # Center alignments for specific parts like the cover page headings
            if "Course Project Final Report" in line or "Team Members" in line or "GitHub Repository" in line:
                align = WD_ALIGN_PARAGRAPH.CENTER
            add_styled_paragraph(doc, line, space_after=8, align=align)
        else:
            # Empty line adds subtle spacing
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
            
        i += 1
        
    print(f"Saving compiled docx to: {report_docx_path}")
    doc.save(report_docx_path)
    print("Compilation completed successfully!")

if __name__ == "__main__":
    main()
